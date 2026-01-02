"""
DOCX 문서 파서

python-docx를 사용하여 DOCX 파일에서 텍스트를 추출하고
페이지 단위로 구조화된 데이터를 생성합니다.
"""

import logging
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.services.document_parser.base_parser import (
    BaseDocumentParser,
    ParsedDocument,
    ParsedPage,
    ParserConfig,
    FileSizeLimitExceededError,
    CorruptedFileError,
)

logger = logging.getLogger(__name__)


class DOCXParser(BaseDocumentParser):
    """DOCX 문서 파서"""

    # 지원하는 파일 확장자
    SUPPORTED_EXTENSIONS = [".docx"]

    # 지원하는 MIME 타입
    SUPPORTED_MIME_TYPES = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def parse(self, file_path: str) -> ParsedDocument:
        """
        DOCX 파일을 파싱하여 구조화된 데이터 반환

        Args:
            file_path: DOCX 파일 경로

        Returns:
            ParsedDocument: 파싱된 문서 데이터

        Raises:
            FileSizeLimitExceededError: 파일 크기 제한 초과
            CorruptedFileError: 손상된 파일
        """
        logger.info(f"DOCX 파싱 시작: {file_path}")

        # Step 1: 파일 존재 여부 확인
        self._validate_file_exists(file_path)

        # Step 2: 파일 타입 검증 [HARD RULE]
        self._validate_file_type(file_path)

        # Step 3: 파일 크기 검증 [HARD RULE]
        self._validate_file_size(file_path)

        # Step 4: DOCX 읽기
        try:
            document = Document(file_path)
        except PackageNotFoundError as e:
            logger.error(f"DOCX 읽기 실패 (패키지 없음): {e}")
            raise CorruptedFileError(f"손상된 DOCX 파일입니다: {e}")
        except Exception as e:
            logger.error(f"예상치 못한 에러: {e}")
            raise CorruptedFileError(f"DOCX 파일을 읽을 수 없습니다: {e}")

        # Step 5: 단락별 텍스트 추출 (페이지 개념 없으므로 단일 페이지로 처리)
        paragraphs_text = []
        for paragraph in document.paragraphs:
            text = paragraph.text
            if self.config.skip_empty_pages and not text.strip():
                continue
            paragraphs_text.append(text)

        # 전체 텍스트 결합
        full_text = "\n".join(paragraphs_text)
        total_characters = len(full_text)

        # Step 6: 메타데이터 추출
        metadata = self._extract_metadata(document)

        # Step 7: ParsedDocument 생성 (DOCX는 페이지 개념이 없으므로 전체를 1페이지로 처리)
        pages = [
            ParsedPage(
                page_number=1,
                content=full_text,
                metadata={
                    "paragraph_count": len(document.paragraphs),
                    "section_count": len(document.sections),
                }
            )
        ]

        result = ParsedDocument(
            pages=pages,
            total_pages=1,
            total_characters=total_characters,
            metadata=metadata,
        )

        logger.info(
            f"DOCX 파싱 완료: {file_path}, "
            f"단락 {len(document.paragraphs)}개, 문자 {total_characters}개"
        )

        return result

    def _extract_metadata(self, document: Document) -> Dict[str, Any]:
        """
        DOCX 메타데이터 추출

        Args:
            document: python-docx Document 객체

        Returns:
            메타데이터 딕셔너리
        """
        metadata = {}

        try:
            core_properties = document.core_properties
            metadata = {
                "title": core_properties.title,
                "author": core_properties.author,
                "subject": core_properties.subject,
                "keywords": core_properties.keywords,
                "comments": core_properties.comments,
                "category": core_properties.category,
                "created": str(core_properties.created) if core_properties.created else None,
                "modified": str(core_properties.modified) if core_properties.modified else None,
                "last_modified_by": core_properties.last_modified_by,
                "revision": core_properties.revision,
            }
            # None 값 제거
            metadata = {k: v for k, v in metadata.items() if v is not None}
        except Exception as e:
            logger.warning(f"메타데이터 추출 실패: {e}")

        return metadata
