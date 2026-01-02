#!/usr/bin/env python3
"""
테스트용 문서 파일 생성 스크립트 (DOCX, TXT, Markdown)

Usage:
    python scripts/generate_test_documents.py
"""

import os
import sys
from pathlib import Path

# 프로젝트 루트를 Python 경로에 추가
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt


def create_output_dirs():
    """출력 디렉토리 생성"""
    base_dir = Path("tests/fixtures")
    docx_dir = base_dir / "docx"
    txt_dir = base_dir / "txt"
    md_dir = base_dir / "markdown"

    docx_dir.mkdir(parents=True, exist_ok=True)
    txt_dir.mkdir(parents=True, exist_ok=True)
    md_dir.mkdir(parents=True, exist_ok=True)

    return docx_dir, txt_dir, md_dir


# ============================================
# DOCX 파일 생성
# ============================================

def generate_sample_valid_docx(output_dir: Path):
    """TC01: 정상 DOCX 파일"""
    file_path = output_dir / "sample_valid.docx"
    doc = Document()

    doc.add_heading("Sample DOCX Document", 0)
    doc.add_paragraph("This is a valid DOCX file for testing purposes.")
    doc.add_paragraph("안녕하세요. 한글 테스트입니다.")

    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("This is the first section with some content.")
    doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit.")

    doc.add_heading("Section 2", level=1)
    doc.add_paragraph("This is the second section.")

    doc.save(str(file_path))
    print(f"✓ Generated: {file_path}")


def generate_sample_empty_docx(output_dir: Path):
    """TC02: 빈 DOCX 파일"""
    file_path = output_dir / "sample_empty.docx"
    doc = Document()
    # 빈 문서
    doc.save(str(file_path))
    print(f"✓ Generated: {file_path}")


def generate_sample_large_docx(output_dir: Path):
    """TC03: 대용량 DOCX 파일 (많은 단락)"""
    file_path = output_dir / "sample_large.docx"
    doc = Document()

    doc.add_heading("Large DOCX Document", 0)

    for i in range(1, 101):  # 100개 섹션
        doc.add_heading(f"Section {i}", level=1)
        for j in range(5):  # 섹션당 5개 단락
            doc.add_paragraph(
                f"This is paragraph {j+1} in section {i}. "
                "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
                * 10  # 긴 단락
            )

    doc.save(str(file_path))
    print(f"✓ Generated: {file_path}")


def generate_sample_korean_docx(output_dir: Path):
    """TC04: 한글 전용 DOCX 파일"""
    file_path = output_dir / "sample_korean.docx"
    doc = Document()

    doc.add_heading("한글 문서 테스트", 0)
    doc.add_paragraph("이것은 한글로 작성된 테스트 문서입니다.")
    doc.add_paragraph("다양한 한글 문장을 포함하고 있습니다.")

    doc.add_heading("첫 번째 섹션", level=1)
    doc.add_paragraph("한글 단락입니다. 가나다라마바사 아자차카타파하.")

    doc.save(str(file_path))
    print(f"✓ Generated: {file_path}")


# ============================================
# TXT 파일 생성
# ============================================

def generate_sample_valid_txt(output_dir: Path):
    """TC01: 정상 TXT 파일"""
    file_path = output_dir / "sample_valid.txt"
    content = """Sample TXT Document
This is a valid text file for testing purposes.
안녕하세요. 한글 테스트입니다.

Section 1
This is the first section with some content.
Lorem ipsum dolor sit amet, consectetur adipiscing elit.

Section 2
This is the second section.
End of document.
"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✓ Generated: {file_path}")


def generate_sample_empty_txt(output_dir: Path):
    """TC02: 빈 TXT 파일"""
    file_path = output_dir / "sample_empty.txt"
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("")
    print(f"✓ Generated: {file_path}")


def generate_sample_large_txt(output_dir: Path):
    """TC03: 대용량 TXT 파일"""
    file_path = output_dir / "sample_large.txt"
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("Large TXT Document\n\n")
        for i in range(1, 1001):  # 1000 라인
            f.write(f"Line {i}: Lorem ipsum dolor sit amet, consectetur adipiscing elit.\n")
    print(f"✓ Generated: {file_path}")


def generate_sample_unicode_txt(output_dir: Path):
    """TC04: 다국어 TXT 파일"""
    file_path = output_dir / "sample_unicode.txt"
    content = """Multilingual Text File

English: Hello World
한국어: 안녕하세요
日本語: こんにちは
中文: 你好
Español: Hola
Français: Bonjour
Deutsch: Hallo
Русский: Здравствуйте
العربية: مرحبا
हिन्दी: नमस्ते
"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✓ Generated: {file_path}")


# ============================================
# Markdown 파일 생성
# ============================================

def generate_sample_valid_md(output_dir: Path):
    """TC01: 정상 Markdown 파일"""
    file_path = output_dir / "sample_valid.md"
    content = """# Sample Markdown Document

This is a valid Markdown file for testing purposes.
안녕하세요. 한글 테스트입니다.

## Section 1

This is the first section with some content.

- List item 1
- List item 2
- List item 3

## Section 2

This is the second section with a code block:

```python
def hello_world():
    print("Hello, World!")
```

### Subsection 2.1

Here's a [link](https://example.com) and an image reference:

![Alt text](image.png)

**Bold text** and *italic text*.

> This is a blockquote.

End of document.
"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✓ Generated: {file_path}")


def generate_sample_empty_md(output_dir: Path):
    """TC02: 빈 Markdown 파일"""
    file_path = output_dir / "sample_empty.md"
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("")
    print(f"✓ Generated: {file_path}")


def generate_sample_readme_md(output_dir: Path):
    """TC03: README 스타일 Markdown"""
    file_path = output_dir / "sample_readme.md"
    content = """# Project Title

[![Build Status](https://travis-ci.org/user/repo.svg?branch=master)](https://travis-ci.org/user/repo)

## Description

This is a sample README file with various Markdown features.

## Installation

```bash
pip install package-name
```

## Usage

```python
from package import module

module.function()
```

## Features

1. Feature one
2. Feature two
3. Feature three

## Contributing

Pull requests are welcome!

## License

MIT
"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✓ Generated: {file_path}")


def generate_sample_korean_md(output_dir: Path):
    """TC04: 한글 Markdown 파일"""
    file_path = output_dir / "sample_korean.md"
    content = """# 한글 마크다운 테스트

이것은 한글로 작성된 마크다운 파일입니다.

## 첫 번째 섹션

한글 내용입니다.

- 리스트 아이템 1
- 리스트 아이템 2
- 리스트 아이템 3

## 두 번째 섹션

코드 블록:

```python
def 안녕():
    print("안녕하세요!")
```

**굵은 글씨**와 *기울임 글씨*

> 인용문입니다.
"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✓ Generated: {file_path}")


def main():
    """메인 함수"""
    print("📄 테스트용 문서 파일 생성 중...\n")

    docx_dir, txt_dir, md_dir = create_output_dirs()

    try:
        # DOCX 파일 생성
        print("\n=== DOCX 파일 생성 ===")
        generate_sample_valid_docx(docx_dir)
        generate_sample_empty_docx(docx_dir)
        generate_sample_large_docx(docx_dir)
        generate_sample_korean_docx(docx_dir)

        # TXT 파일 생성
        print("\n=== TXT 파일 생성 ===")
        generate_sample_valid_txt(txt_dir)
        generate_sample_empty_txt(txt_dir)
        generate_sample_large_txt(txt_dir)
        generate_sample_unicode_txt(txt_dir)

        # Markdown 파일 생성
        print("\n=== Markdown 파일 생성 ===")
        generate_sample_valid_md(md_dir)
        generate_sample_empty_md(md_dir)
        generate_sample_readme_md(md_dir)
        generate_sample_korean_md(md_dir)

        print("\n✅ 모든 문서 파일 생성 완료!")
        print(f"📁 위치:")
        print(f"   - DOCX: {docx_dir.absolute()}")
        print(f"   - TXT: {txt_dir.absolute()}")
        print(f"   - Markdown: {md_dir.absolute()}")

    except Exception as e:
        print(f"\n❌ 에러 발생: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())
